#app.py

import pandas as pd
import io
import os
import jwt
import datetime
import base64
from models import OriginalFile, TranslatedFile, User
from werkzeug.utils import secure_filename
from googletrans import Translator
from docx import Document
from flask_cors import cross_origin
from io import BytesIO
from database import *
from PyPDF2 import PdfReader
from flask import Flask, request, jsonify
from flask_cors import CORS

app = Flask(__name__)
CORS(app, origins=['http://localhost:3000', 'http://localhost:3001'])
translator = Translator()

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your_secret_key')

def create_token(user_id):
    token = jwt.encode(
        {
            'user_id': user_id,
            'exp': datetime.datetime.now() + datetime.timedelta(hours=24)
        },
        app.config['SECRET_KEY'],
        algorithm='HS256'
    )
    print("This the token " + token)
    return token

@app.route('/signup', methods=['POST'])
def signup():
    data = request.json
    first_name = data.get('first_name')
    last_name = data.get('last_name')
    email = data.get('email')
    credential = data.get('credential')
        
    add_user(first_name, last_name, email, credential)
    
    user = session.query(User).filter_by(email=email).first()
    token = create_token(user.user_id)
    
    return jsonify({'token': token}), 201

@app.route('/signin', methods=['POST'])
def signin():
    try:
        data = request.json
        email = data.get('email')
        credential = data.get('credential')
        
        response = validate_user(email, credential)
        if(response == 1):
            return jsonify({'message': 'Invalid credentials'}), 401
        elif response == 2:
            return jsonify({'message': 'User not found'}), 404
        else: 
            user = session.query(User).filter_by(email=email).first()
            token = create_token(user.user_id)
            return jsonify({'token': token}), 200
    except Exception as e:
        session.rollback()  # Rollback the transaction if an error occurs
        print("Error:", e)
        return jsonify({"status": "error", "message": "An error occurred"})
    
@app.route('/emailcheck', methods=['POST'])
@cross_origin(origin='localhost',headers=['Content- Type','Authorization'])
def emailcheck():
    email = request.json.get('email')
    email_exists = check_for_email(email)
    return jsonify({'exists': email_exists})

@app.route('/getname', methods=['POST'])
def getname():
    userId = request.json.get('userId')
    user = session.query(User).filter_by(user_id=userId).first()
    if user:
        user_data = {'name': user.first_name}
        return jsonify(user_data)
    else:
        return jsonify("Idk bro the user ID can't be found")

@app.route('/translate', methods=['POST'])
def translate():
    file = request.files['file']
    src_lang = request.form.get('srcLang')
    dest_lang = request.form.get('destLang')
    translated_text = None
    doc_b64 = None
    if file:
        if file.filename.endswith('.docx'):
            doc = Document(io.BytesIO(file.read()))
            translated_paragraphs = []
            for index, para in enumerate(doc.paragraphs):
                translated_text = translator.translate(para.text, dest=dest_lang, src=src_lang).text
                translated_paragraphs.append(translated_text)
                doc.paragraphs[index].text = translated_text
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            doc_b64 = base64.b64encode(doc_io.read()).decode()
        elif file.filename.endswith('.txt'):
            text_to_translate = file.read().decode('utf-8')
            translated_text = translator.translate(text_to_translate, dest=dest_lang, src=src_lang).text
            with open('translated.txt', 'w', encoding='utf-8') as f:
                f.write(translated_text)
            with open('translated.txt', 'rb') as f:
                doc_b64 = base64.b64encode(f.read()).decode()
        elif file.filename.endswith('.pdf'):
            pdf = PdfReader(file)
            text_to_translate = ' '.join([page.extract_text() for page in pdf.pages])
            translated_text = translator.translate(text_to_translate, dest=dest_lang, src=src_lang).text
            doc_b64 = base64.b64encode(translated_text.encode('utf-8')).decode()
        elif file.filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file)
            text_columns = df.select_dtypes(include=[object]).columns
            translated_df = df.copy()
            for col in text_columns:
                translated_df[col] = df[col].apply(lambda cell: translator.translate(str(cell), dest=dest_lang, src=src_lang).text.title())
            translated_df.columns = [translator.translate(col, dest=dest_lang, src=src_lang).text for col in translated_df.columns]
            translated_text = translated_df.to_csv(index=False)
            doc_b64 = base64.b64encode(translated_text.encode('utf-8')).decode()
            
        return jsonify(translated_text=translated_text, translated_file=doc_b64)
    else:
        return jsonify(error="No file provided"), 400

@app.route('/savefiles', methods=['POST'])
def savefiles():
    data = request.json
    
@app.route('/upload', methods=['POST'])
def upload():
    try:
        og_file = request.files['ogfile']
        og_file_name = og_file.filename
        src_lang = request.form.get('srcLang')
        dest_lang = request.form.get('destLang')
        trans_file = request.files['transFile']
        user_id = request.form.get('userId')

        # Check if the original file already exists
        file_already = session.query(OriginalFile).filter_by(file_name=og_file_name, user_id=user_id).first()
        
        if file_already:
            og_file_id = file_already.file_id
        else:
            # Add the original file if it does not exist
            add_original_file(user_id, og_file, src_lang, datetime.datetime.now())
            file_already = session.query(OriginalFile).filter_by(file_name=og_file_name, user_id=user_id).first()
            if file_already:
                og_file_id = file_already.file_id
            else:
                return jsonify(status='error', message='Failed to add original file'), 500
        
        # Add the translated file
        success = add_trans_file(og_file_id, trans_file, dest_lang, datetime.datetime.now())
        return jsonify(status = 'Status', message=success);
    
    except Exception as e:
        session.rollback()  # Rollback the transaction if an error occurs
        print("Error:", e)
        return jsonify(status='error', message=str(e)), 500


@app.route('/populate', methods=['POST'])
def populate():
    try:
        request_data = request.json
        user_id = request_data.get('userId')    
        
        table_data = get_users_files(user_id)   
        return jsonify(table_data), 200
    except Exception as e:
        session.rollback()  # Rollback the transaction if an error occurs
        print("Error:", e)
        return jsonify(status='error', message=str(e)), 500

@app.route('/download', methods=['POST'])
def download():
    try:
        request_data = request.json
        user_id = request_data.get('userId')
        og_file_name = request_data.get('key') 
        trans_lan = request_data.get('item')   
        
        return download_file(user_id, og_file_name, trans_lan);
    except Exception as e:
        session.rollback()  # Rollback the transaction if an error occurs
        print("Error:", e)
        return jsonify(status='error', message=str(e)), 500

@app.route('/tobetranslated', methods=['POST'])
def tobetranslated():
    try:
        request_data = request.json
        user_id = request_data.get('userId')
        
        files = get_users_og_files(user_id);
        return files
    except Exception as e:
        session.rollback()  # Rollback the transaction if an error occurs
        print("Error:", e)
        return jsonify(status='error', message=str(e)), 500
   
if __name__ == '__main__':
    app.run(debug=True)